"""Markdown → docx 报告生成服务（V1.1）。

将报告总结功能的最终 Markdown 内容按「通用报告格式」渲染为 Word 文档：

- 报告标题：黑体 三号(16pt) 居中
- 一级标题（##）：黑体 四号(14pt)
- 二级标题（###）：宋体 小四(12pt) 加粗
- 正文：宋体 小四(12pt)，首行缩进 2 字符，1.5 倍行距
- 页脚：居中页码（PAGE 域）
- 图片：解析 Markdown 图片引用，从磁盘读取并居中插入（页宽内缩放）
- 表格：Markdown 表格转带边框的 Word 表格

图片 URL → 磁盘路径的解析规则：
- /api/uploads/{kb_id}/{path}     → {UPLOAD_DIR}/{kb_id}/{path}（知识库图片）
- /api/report/uploads/{sid}/{file} → {UPLOAD_DIR}/report/{sid}/{file}（报告上传图片）
"""
import os
import re
import logging

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import settings

logger = logging.getLogger("app.docx")

# ---------------- 字号常量（中文字号 → pt） ----------------
TITLE_PT = 16       # 三号
H1_PT = 14          # 四号
H2_PT = 12          # 小四
BODY_PT = 12        # 小四
TABLE_PT = 10.5     # 五号（表格略小，防止溢出）
CAPTION_PT = 10.5

FONT_HEI = "黑体"     # SimHei
FONT_SONG = "宋体"    # SimSun

# Markdown 内联加粗 **xx**
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
# 图片引用 ![alt](src)，src 可含括号
_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
# 引用行 > text
_QUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
# 无序列表 - / * / +
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
# 有序列表 1. text
_ORDERED_RE = re.compile(r"^\s*(\d+)[.、)]\s+(.*)$")


def _set_run_font(run, font_cn: str, size_pt: float, bold: bool = False, italic: bool = False) -> None:
    """同时设置西文与中文字体（eastAsia），否则 Word 中中文显示为默认西文字体。"""
    run.font.name = font_cn
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_cn)
    rFonts.set(qn("w:hAnsi"), font_cn)


def _add_paragraph(doc, text: str, font_cn: str = FONT_SONG, size_pt: float = BODY_PT,
                   bold: bool = False, indent_chars: float = 2.0, line_spacing: float = 1.5,
                   align: int = WD_ALIGN_PARAGRAPH.LEFT, space_before: float = 0, space_after: float = 0):
    """新增正文/标题段落，支持行内加粗。"""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_chars:
        # 中文字符缩进按字号计算：2 字符 = 2 * 字号
        pf.first_line_indent = Pt(size_pt * indent_chars)
    # 行内加粗解析：**text** 交替分段
    parts = _BOLD_RE.split(text or "")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        _set_run_font(run, font_cn, size_pt, bold=(bold or i % 2 == 1))
    return p


def _add_page_number_footer(doc) -> None:
    """页脚添加居中页码域（PAGE）。"""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    _set_run_font(run, FONT_SONG, 9)


def _setup_section(doc) -> None:
    """A4 页面、标准页边距。"""
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)


def _resolve_image_path(url: str) -> str | None:
    """报告 Markdown 中的图片 URL → 磁盘绝对路径。无法解析返回 None。"""
    from urllib.parse import urlsplit, unquote
    path = urlsplit(url or "").path
    path = unquote(path)
    if path.startswith("/api/report/uploads/"):
        rel = path[len("/api/report/uploads/"):]
        return os.path.join(settings.UPLOAD_DIR, "report", rel)
    if path.startswith("/api/uploads/"):
        rel = path[len("/api/uploads/"):]
        return os.path.join(settings.UPLOAD_DIR, rel)
    return None


def _add_image(doc, url: str, alt: str) -> None:
    """插入图片段落（居中，宽度不超过 14cm）。缺失时输出文字占位。"""
    abs_path = _resolve_image_path(url)
    if abs_path and os.path.isfile(abs_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run()
            # 读取原始像素宽，仅对较小的图片保持原尺寸，其余（含无法读取尺寸时）按页宽缩放
            try:
                from PIL import Image
                with Image.open(abs_path) as im:
                    w_px, _ = im.size
            except Exception:
                w_px = 0
            if w_px > 0 and w_px <= 560:  # ~14cm @ 96dpi 以内的保持原图
                run.add_picture(abs_path)
            else:
                run.add_picture(abs_path, width=Cm(14))
            if alt:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_before = Pt(0)
                cap.paragraph_format.space_after = Pt(6)
                c_run = cap.add_run(alt)
                _set_run_font(c_run, FONT_SONG, CAPTION_PT)
        except Exception as e:
            logger.warning("插入图片失败 %s: %s", url, e)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[图片：{alt or url}]")
        _set_run_font(run, FONT_SONG, BODY_PT)


def _add_table(doc, rows: list[list[str]]) -> None:
    """Markdown 表格 → 带边框的 Word 表格。rows 不含分隔行。"""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            text = row[ci].strip() if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(text)
            _set_run_font(run, FONT_SONG, TABLE_PT, bold=(ri == 0))
    # 表后留一点间距
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)
    sp_run = sp.add_run("")
    _set_run_font(sp_run, FONT_SONG, 2)


def _render_block(doc, block: dict) -> None:
    """渲染一个解析后的块。"""
    kind = block["type"]
    if kind == "heading1":
        _add_paragraph(doc, block["text"], FONT_HEI, H1_PT, indent_chars=0,
                       space_before=12, space_after=6)
    elif kind == "heading2":
        _add_paragraph(doc, block["text"], FONT_SONG, H2_PT, bold=True, indent_chars=0,
                       space_before=8, space_after=4)
    elif kind == "title":
        _add_paragraph(doc, block["text"], FONT_HEI, TITLE_PT, indent_chars=0,
                       align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=14)
    elif kind == "image":
        _add_image(doc, block["url"], block["alt"])
    elif kind == "table":
        _add_table(doc, block["rows"])
    elif kind == "bullet":
        _add_paragraph(doc, "• " + block["text"], indent_chars=0)
    elif kind == "ordered":
        _add_paragraph(doc, f"{block['num']}. " + block["text"], indent_chars=0)
    elif kind == "quote":
        _add_paragraph(doc, block["text"], indent_chars=0)
    elif kind == "code":
        _add_paragraph(doc, block["text"], indent_chars=0)
    else:  # paragraph
        _add_paragraph(doc, block["text"])


def _parse_blocks(markdown: str) -> list[dict]:
    """将 Markdown 逐行解析为块列表。"""
    blocks: list[dict] = []
    lines = (markdown or "").split("\n")
    i = 0
    n = len(lines)
    ordered_num = 0

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            ordered_num = 0
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            blocks.append({"type": "heading1" if level <= 2 else "heading2", "text": text})
            i += 1
            continue

        # 代码块 ``` 或 ~~~
        if stripped.startswith("```") or stripped.startswith("~~~"):
            fence = stripped[:3]
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith(fence):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束围栏
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        # 图片独立行
        img = _IMAGE_RE.match(stripped)
        if img and not stripped.startswith("|"):
            blocks.append({"type": "image", "alt": img.group(1).strip(), "url": img.group(2).strip()})
            i += 1
            continue

        # 表格块：连续以 | 开头结尾的行
        if stripped.startswith("|") and stripped.endswith("|"):
            rows: list[list[str]] = []
            while i < n:
                s = lines[i].strip()
                if not (s.startswith("|") and s.endswith("|")):
                    break
                cells = [c.strip() for c in s.strip("|").split("|")]
                # 跳过 |---|---| 分隔行
                if all(re.match(r"^[-:\s]+$", c) for c in cells):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": rows})
            ordered_num = 0
            continue

        # 引用
        m = _QUOTE_RE.match(stripped)
        if m:
            blocks.append({"type": "quote", "text": m.group(1)})
            i += 1
            continue

        # 无序列表
        m = _BULLET_RE.match(stripped)
        if m:
            ordered_num = 0
            blocks.append({"type": "bullet", "text": m.group(1)})
            i += 1
            continue

        # 有序列表
        m = _ORDERED_RE.match(stripped)
        if m:
            ordered_num += 1
            blocks.append({"type": "ordered", "num": int(m.group(1)) if m.group(1).isdigit() else ordered_num,
                           "text": m.group(2)})
            i += 1
            continue

        # 普通段落（含行内图片语法，仅文本展示链接来源）
        blocks.append({"type": "paragraph", "text": stripped})
        i += 1

    return blocks


def generate_docx(content_markdown: str, title: str, out_path: str) -> str:
    """生成 docx 报告，返回文件路径。

    Args:
        content_markdown: 报告正文 Markdown（由大模型生成）
        title: 报告标题
        out_path: 输出文件路径
    """
    doc = Document()
    # 默认样式兜底
    normal = doc.styles["Normal"]
    normal.font.name = FONT_SONG
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    normal.font.size = Pt(BODY_PT)

    _setup_section(doc)
    _add_page_number_footer(doc)

    # 报告标题
    blocks = [{"type": "title", "text": title.strip() or "报告总结"}]
    blocks.extend(_parse_blocks(content_markdown))

    for block in blocks:
        _render_block(doc, block)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    logger.info("docx 报告已生成: %s", out_path)
    return out_path
