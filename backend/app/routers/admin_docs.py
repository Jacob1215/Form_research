"""管理后台 — 文档上传/列表/详情/删除/下载/解析。

支持 PDF 文档解析：提取文本、表格、图片，解析后可在预览框中查看和搜索。
"""
import os
import json
import logging

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy import select, func
from sqlalchemy.orm import Session

from ..database import get_db
from ..config import settings
from ..models import KnowledgeBase, Document
from ..schemas import DocumentOut, DocumentDetail
from ..image_utils import rewrite_md_image_paths

logger = logging.getLogger("app.admin_docs")

router = APIRouter(prefix="/api/admin", tags=["admin-docs"])

ALLOWED_EXT = {"pdf", "doc", "docx", "txt", "md"}
TEXT_EXT = {"txt", "md"}
MAX_SIZE_BYTES = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024


def _ext(filename: str) -> str:
    return os.path.splitext(filename)[1].lower().lstrip(".")


def _cap_text(text: str, file_path: str) -> str:
    """按 DOC_TEXT_CAP 截断入库文本，并在截断时记告警。

    V1.2.3：大规范（如公路隧道设计规范）远超 10 万字符，硬截断会丢掉靠后章节/表格，
    表现为"查不到靠后内容"。上限改为可配置，且截断时明确告警便于排查。
    """
    if len(text) > settings.DOC_TEXT_CAP:
        logger.warning(
            "文档 %s 文本 %d 字符超过上限 %d，已截断（靠后内容不入库）",
            file_path, len(text), settings.DOC_TEXT_CAP,
        )
    return text[: settings.DOC_TEXT_CAP]


def _extract_text(file_path: str, file_type: str) -> str | None:
    """从文档中提取文本内容，供 RAG 检索使用。

    支持：txt/md（直接读取）、pdf（pdfplumber，中文效果好）、docx（python-docx）。
    对扫描版 PDF 自动降级为 PyMuPDF 文本提取。
    """
    try:
        if file_type in TEXT_EXT:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return _cap_text(text, file_path)

        if file_type == "pdf":
            # 优先使用 pdfplumber（中文支持远优于 pypdf）
            import pdfplumber
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            full_text = "\n\n".join(parts) if parts else ""

            # 如果 pdfplumber 提取内容过少，可能是扫描件，尝试 PyMuPDF
            if len(full_text.strip()) < 200:
                logger.info("PDF 可能为扫描件，尝试 PyMuPDF 提取: %s", file_path)
                try:
                    import fitz
                    fitz_doc = fitz.open(file_path)
                    ocr_parts = []
                    for page in fitz_doc:
                        text = page.get_text()
                        if text and len(text.strip()) > 50:
                            ocr_parts.append(text)
                    fitz_doc.close()
                    if ocr_parts:
                        full_text = "\n\n".join(ocr_parts)
                        logger.info("PyMuPDF 提取成功: %d 页有文本", len(ocr_parts))
                except Exception as e:
                    logger.warning("PyMuPDF 提取失败: %s", e)

            return _cap_text(full_text, file_path) if full_text.strip() else None

        if file_type == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return _cap_text("\n".join(parts), file_path) if parts else None

        return None
    except Exception as e:  # noqa: BLE001
        logger.warning("提取文本失败 %s: %s", file_path, e)
        return None


def _parse_pdf(file_path: str, doc_id: int, kb_id: int) -> dict:
    """解析 PDF 文件，提取文本、表格和图片。

    返回结构：
    {
        "pages": [
            {"page_num": 1, "text": "...", "tables": [[...]], "images": [{"id": "...", "src": "..."}]}
        ],
        "total_pages": N
    }
    """
    import pdfplumber
    import fitz  # PyMuPDF

    # 图片存储目录
    img_dir = os.path.join(settings.UPLOAD_DIR, str(kb_id), "images", str(doc_id))
    os.makedirs(img_dir, exist_ok=True)

    pages = []
    full_text_parts = []

    # 使用 pdfplumber 提取文本和表格
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []

            # 清理表格中的 None 值
            cleaned_tables = []
            for table in tables:
                cleaned = [[(cell or "").strip() for cell in row] for row in table]
                cleaned_tables.append(cleaned)

            pages.append({
                "page_num": i,
                "text": text,
                "tables": cleaned_tables,
                "images": [],  # 先占位，后面用 PyMuPDF 填充
            })
            full_text_parts.append(text)

    # 使用 PyMuPDF 提取图片
    fitz_doc = fitz.open(file_path)
    try:
        for page_num in range(len(fitz_doc)):
            page = fitz_doc[page_num]
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = fitz_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # 跳过过小的图片（< 1KB，可能是装饰元素）
                    if len(image_bytes) < 1024:
                        continue

                    img_filename = f"p{page_num + 1}_img{img_index}.{image_ext}"
                    img_path = os.path.join(img_dir, img_filename)
                    with open(img_path, "wb") as f:
                        f.write(image_bytes)

                    img_id = f"p{page_num + 1}_img{img_index}"
                    img_src = f"/api/admin/documents/{doc_id}/images/{img_filename}"

                    if page_num < len(pages):
                        pages[page_num]["images"].append({
                            "id": img_id,
                            "src": img_src,
                        })
                except Exception as e:  # noqa: BLE001
                    logger.warning("提取图片失败 p%d img%d: %s", page_num + 1, img_index, e)
    finally:
        fitz_doc.close()

    return {
        "pages": pages,
        "total_pages": len(pages),
    }


def _parse_markdown(file_path: str) -> dict:
    """解析 Markdown 文件，提取文本、表格和图片引用。

    按标题（## ）分页，提取 Markdown 表格和图片引用。
    返回结构与 _parse_pdf 一致。
    """
    import re as _re

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    lines = content.split("\n")
    pages: list[dict] = []
    page_num = 0

    current_text: list[str] = []
    current_tables: list[list[list[str]]] = []
    current_images: list[dict] = []
    current_table: list[list[str]] = []
    in_table = False

    def _flush_page():
        nonlocal page_num
        if not current_text and not current_tables and not current_images:
            return
        page_num += 1
        if current_table:
            current_tables.append(current_table)
        pages.append({
            "page_num": page_num,
            "text": "\n".join(current_text).strip(),
            "tables": current_tables,
            "images": current_images,
        })

    for line in lines:
        stripped = line.strip()

        # 标题行 → 新页
        if _re.match(r'^#{1,6}\s+', stripped):
            _flush_page()
            current_text = [stripped]
            current_tables = []
            current_images = []
            current_table = []
            in_table = False
            continue

        # Markdown 表格行
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # 跳过分隔行 |---|---|
            if all(_re.match(r'^[-:\s]+$', c) for c in cells):
                in_table = True
                continue
            current_table.append(cells)
            in_table = True
            continue

        # 表格结束
        if in_table and not (stripped.startswith("|") and stripped.endswith("|")):
            if current_table:
                current_tables.append(current_table)
            current_table = []
            in_table = False

        # 图片引用 ![alt](src)
        img_match = _re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_src = img_match.group(2)
            current_images.append({
                "id": f"img_{len(current_images)}",
                "src": img_src,
            })
            current_text.append(alt_text or "[图片]")
            continue

        current_text.append(line)

    # 最后一页
    _flush_page()

    # 如果没有任何标题，整篇作为一个页面
    if not pages:
        pages.append({
            "page_num": 1,
            "text": content,
            "tables": [],
            "images": [],
        })

    return {"pages": pages, "total_pages": len(pages)}


# ---------- 文档列表 ----------

@router.get("/knowledge-bases/{kb_id}/documents")
def list_documents(kb_id: int, db: Session = Depends(get_db)):
    kb = db.get(KnowledgeBase, kb_id)
    if kb is None:
        raise HTTPException(status_code=404, detail="知识库不存在")
    rows = db.execute(
        select(Document).where(Document.kb_id == kb_id).order_by(Document.created_at.desc())
    ).scalars().all()
    return {"items": [DocumentOut.from_orm(r).model_dump() for r in rows]}


# ---------- 上传 ----------

@router.post("/knowledge-bases/{kb_id}/documents")
def upload_documents(
    kb_id: int,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    kb = db.get(KnowledgeBase, kb_id)
    if kb is None:
        raise HTTPException(status_code=404, detail="知识库不存在")
    if not files:
        raise HTTPException(status_code=400, detail="未提供文件")
    if len(files) > settings.MAX_UPLOAD_FILES:
        raise HTTPException(status_code=400, detail=f"单次最多上传 {settings.MAX_UPLOAD_FILES} 个文件")

    kb_dir = os.path.join(settings.UPLOAD_DIR, str(kb_id))
    os.makedirs(kb_dir, exist_ok=True)

    created: list[Document] = []
    for upload in files:
        ext = _ext(upload.filename or "")
        if ext not in ALLOWED_EXT:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型：{upload.filename}")
        data = upload.file.read()
        if len(data) > MAX_SIZE_BYTES:
            raise HTTPException(status_code=400, detail=f"文件过大（>{settings.MAX_UPLOAD_SIZE_MB}MB）：{upload.filename}")
        base_name = os.path.splitext(upload.filename or "uploaded")[0]
        unique_name = f"{base_name}_{os.urandom(3).hex()}.{ext}"
        save_path = os.path.join(kb_dir, unique_name)
        with open(save_path, "wb") as f:
            f.write(data)

        # 仅对 txt/md 提取文本内容，供检索使用
        content_text = _extract_text(save_path, ext)

        doc = Document(
            kb_id=kb_id,
            file_name=upload.filename or unique_name,
            file_type=ext,
            file_size=len(data),
            file_path=save_path,
            content_text=content_text,
        )
        db.add(doc)
        db.flush()
        created.append(doc)

    db.commit()
    for d in created:
        db.refresh(d)

    _refresh_kb_doc_count(db, kb_id)

    # 自动索引：对上传的文档进行分块和向量化
    for d in created:
        try:
            _auto_index_document(db, d)
        except Exception as e:
            logger.warning("自动索引文档 %d 失败: %s", d.id, e)

    return {"items": [DocumentOut.from_orm(d).model_dump() for d in created]}


# ---------- 文件夹上传 ----------

# 文件夹上传支持的文件类型（md 创建 Document；图片直接落地）
FOLDER_DOC_EXT = {"md", "markdown"}
FOLDER_IMG_EXT = {"png", "jpg", "jpeg", "gif", "webp", "svg", "bmp"}
# 过滤掉的系统文件
_IGNORE_FILES = {".ds_store", "thumbs.db", "desktop.ini", "__macosx"}


@router.post("/knowledge-bases/{kb_id}/documents/folder")
def upload_folder(
    kb_id: int,
    files: list[UploadFile] = File(...),
    relative_paths: str | None = Form(None),
    db: Session = Depends(get_db),
):
    """文件夹批量上传：保留相对目录结构。

    - .md/.markdown 文件：创建 Document 记录，content_text 直接读取 md 全文
    - 图片文件：存到对应目录（不建 Document），供 md 预览时引用
    - 其他文件：忽略
    - 系统文件（.DS_Store 等）：忽略

    前端通过 webkitdirectory 选文件夹，file.webkitRelativePath 携带相对路径，
    通过 relative_paths（JSON 数组）字段传递，与 files 按下标一一对应。
    """
    kb = db.get(KnowledgeBase, kb_id)
    if kb is None:
        raise HTTPException(status_code=404, detail="知识库不存在")
    if not files:
        raise HTTPException(status_code=400, detail="未提供文件")

    # 解析前端传入的相对路径列表（与 files 按下标对齐）
    paths: list[str] = []
    if relative_paths:
        try:
            raw = json.loads(relative_paths)
            if isinstance(raw, list):
                paths = [str(p) for p in raw]
        except (json.JSONDecodeError, ValueError):
            logger.warning("relative_paths 解析失败，回退使用文件名: %s", relative_paths[:200])

    kb_dir = os.path.join(settings.UPLOAD_DIR, str(kb_id))
    os.makedirs(kb_dir, exist_ok=True)

    created: list[Document] = []
    img_count = 0
    skipped = 0

    for i, upload in enumerate(files):
        filename = upload.filename or ""
        # 优先使用前端传入的相对路径（保留目录结构）；缺失/为空时回退到 upload.filename（兼容旧前端）
        relative_path = paths[i] if i < len(paths) and paths[i] else filename
        if not relative_path:
            relative_path = os.path.basename(filename)

        ext = _ext(filename)
        low_name = filename.lower()

        # 跳过系统文件
        if any(ign in low_name for ign in _IGNORE_FILES):
            skipped += 1
            continue

        data = upload.file.read()
        if len(data) > MAX_SIZE_BYTES:
            logger.warning("文件过大（>%dMB），跳过: %s", settings.MAX_UPLOAD_SIZE_MB, filename)
            skipped += 1
            continue

        # 构造存储路径：保留相对目录结构
        # 标准化路径分隔符
        rel_path = relative_path.replace("\\", "/").lstrip("./")
        save_path = os.path.join(kb_dir, rel_path)
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        with open(save_path, "wb") as f:
            f.write(data)

        if ext in FOLDER_DOC_EXT:
            # md 文件：创建 Document
            content_text = _cap_text(data.decode("utf-8", errors="ignore"), save_path)
            doc = Document(
                kb_id=kb_id,
                file_name=os.path.basename(filename),
                file_type="md",
                file_size=len(data),
                file_path=save_path,
                relative_path=rel_path,
                content_text=content_text,
                parse_status="done",
                parsed_content=json.dumps(
                    {"pages": [{"page_num": 1, "text": content_text, "tables": [], "images": []}],
                     "total_pages": 1},
                    ensure_ascii=False,
                ),
            )
            db.add(doc)
            db.flush()
            created.append(doc)
        elif ext in FOLDER_IMG_EXT:
            img_count += 1
        else:
            skipped += 1

    db.commit()
    for d in created:
        db.refresh(d)

    _refresh_kb_doc_count(db, kb_id)

    # 自动索引 md 文档
    for d in created:
        try:
            _auto_index_document(db, d)
        except Exception as e:
            logger.warning("自动索引文档 %d 失败: %s", d.id, e)

    logger.info(
        "文件夹上传完成 kb_id=%d: %d 个 md 文档, %d 张图片, %d 个跳过",
        kb_id, len(created), img_count, skipped,
    )

    return {
        "items": [DocumentOut.from_orm(d).model_dump() for d in created],
        "image_count": img_count,
        "skipped": skipped,
    }


# ---------- 文档详情 ----------

@router.get("/documents/{doc_id}")
def get_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 对 md 文件：返回重写图片路径后的 content_text（供前端 ReactMarkdown 直接渲染）
    display_text = doc.content_text
    if doc.file_type in ("md", "markdown") and display_text:
        display_text = rewrite_md_image_paths(display_text, doc.kb_id, doc.file_path)

    return DocumentDetail(
        id=doc.id,
        kb_id=doc.kb_id,
        file_name=doc.file_name,
        file_type=doc.file_type,
        file_size=doc.file_size,
        content_text=display_text,
        parsed_content=doc.parsed_content,
        parse_status=getattr(doc, "parse_status", "pending") or "pending",
        created_at=doc.created_at.isoformat() if doc.created_at else None,
        updated_at=doc.updated_at.isoformat() if doc.updated_at else None,
    )


# ---------- 文档解析 ----------

@router.post("/documents/{doc_id}/parse")
def parse_document(doc_id: int, db: Session = Depends(get_db)):
    """解析文档：支持 PDF 和 Markdown 文件，提取文本、表格、图片。"""
    doc = db.get(Document, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="文档不存在")

    if doc.file_type not in ("pdf", "md", "markdown"):
        raise HTTPException(status_code=400, detail="仅支持 PDF 和 Markdown 文件解析")

    if not doc.file_path or not os.path.isfile(doc.file_path):
        raise HTTPException(status_code=404, detail="文件不存在")

    # 更新状态为解析中
    doc.parse_status = "parsing"
    db.commit()

    try:
        if doc.file_type == "pdf":
            result = _parse_pdf(doc.file_path, doc.id, doc.kb_id)
        else:
            result = _parse_markdown(doc.file_path)

        # 存储解析结果
        doc.parsed_content = json.dumps(result, ensure_ascii=False)
        doc.parse_status = "done"

        # 同时更新 content_text（用于 RAG 检索）
        full_text = "\n\n".join(p["text"] for p in result["pages"] if p.get("text"))
        if full_text:
            doc.content_text = _cap_text(full_text, doc.file_path)

        db.commit()
        db.refresh(doc)

        # 解析后 content_text 已更新，必须重建分块；否则混合检索仍使用上传时基于旧文本建的分块，
        # 导致“解析后仍查不到规范内容”的问题（V1.0.7 修复）
        try:
            rebuilt = _auto_index_document(db, doc)
            logger.info("文档 %s 解析后重建分块：%d 个", doc_id, rebuilt)
        except Exception as e:  # noqa: BLE001
            logger.warning("文档 %s 解析后重建分块失败: %s", doc_id, e)

        logger.info("文档 %s 解析完成：%d 页，%d 张图片",
                     doc_id, result["total_pages"],
                     sum(len(p.get("images", [])) for p in result["pages"]))

        return {
            "success": True,
            "parse_status": "done",
            "total_pages": result["total_pages"],
            "total_images": sum(len(p.get("images", [])) for p in result["pages"]),
            "total_tables": sum(len(p.get("tables", [])) for p in result["pages"]),
        }

    except Exception as e:
        doc.parse_status = "error"
        db.commit()
        logger.error("文档 %s 解析失败: %s", doc_id, e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"解析失败: {e}")


# ---------- 图片服务 ----------

@router.get("/documents/{doc_id}/images/{img_filename}")
def serve_document_image(doc_id: int, img_filename: str, db: Session = Depends(get_db)):
    """提供解析后的图片文件。"""
    doc = db.get(Document, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="文档不存在")

    img_path = os.path.join(settings.UPLOAD_DIR, str(doc.kb_id), "images", str(doc_id), img_filename)
    if not os.path.isfile(img_path):
        raise HTTPException(status_code=404, detail="图片不存在")

    return FileResponse(img_path)


# ---------- 删除 ----------

@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="文档不存在")
    try:
        if doc.file_path and os.path.isfile(doc.file_path):
            os.remove(doc.file_path)
    except OSError as e:
        logger.warning("删除文件失败 %s: %s", doc.file_path, e)
    kb_id = doc.kb_id

    # 清理向量分块
    try:
        from ..vector_store import VectorStore
        VectorStore(db).delete_chunks(doc_id)
    except Exception as e:
        logger.warning("清理分块失败 doc_id=%d: %s", doc_id, e)

    db.delete(doc)
    db.commit()
    _refresh_kb_doc_count(db, kb_id)
    return {"success": True}


# ---------- 下载 ----------

@router.get("/documents/{doc_id}/download")
def download_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="文档不存在")
    if not doc.file_path or not os.path.isfile(doc.file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(doc.file_path, filename=doc.file_name)


# ---------- 重建索引 ----------

@router.post("/documents/{doc_id}/reindex")
def reindex_document(doc_id: int, db: Session = Depends(get_db)):
    """手动重建文档的向量索引：分块 → 向量化 → 存储。"""
    doc = db.get(Document, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="文档不存在")
    if not doc.content_text:
        raise HTTPException(status_code=400, detail="文档需先有文本内容，请先上传或解析文档")

    try:
        result = _auto_index_document(db, doc)
        return {"success": True, "chunks": result}
    except Exception as e:
        logger.error("重建索引失败 doc_id=%d: %s", doc_id, e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"索引重建失败: {e}")


@router.post("/knowledge-bases/{kb_id}/reindex-all")
def reindex_knowledge_base(kb_id: int, db: Session = Depends(get_db)):
    """重建知识库中所有文档的向量索引。"""
    kb = db.get(KnowledgeBase, kb_id)
    if kb is None:
        raise HTTPException(status_code=404, detail="知识库不存在")

    docs = db.execute(
        select(Document).where(
            Document.kb_id == kb_id,
            Document.content_text.isnot(None),
        )
    ).scalars().all()

    total_chunks = 0
    errors = []
    for doc in docs:
        try:
            total_chunks += _auto_index_document(db, doc)
        except Exception as e:
            errors.append({"doc_id": doc.id, "file_name": doc.file_name, "error": str(e)})
            logger.warning("重建索引失败 doc_id=%d: %s", doc.id, e)

    return {
        "success": True,
        "total_docs": len(docs),
        "total_chunks": total_chunks,
        "errors": errors,
    }


# ---------- 内部工具 ----------

def _auto_index_document(db: Session, doc: Document) -> int:
    """对文档进行分块、向量化并存储到 pgvector。

    返回创建的分块数。如果文档没有文本内容，返回 0。
    """
    if not doc.content_text:
        return 0

    from ..chunking_service import ChunkingService
    from ..embedding_service import embedding_service
    from ..vector_store import VectorStore

    # 1. 分块
    chunker = ChunkingService()
    chunks = chunker.chunk_text(doc.content_text, doc.file_name)
    if not chunks:
        logger.warning("文档 %d 分块结果为空", doc.id)
        return 0

    # 2. 向量化
    chunk_texts = [c.content for c in chunks]
    embeddings = None
    try:
        embeddings = embedding_service.embed(chunk_texts)
    except Exception as e:
        logger.warning("文档 %d 向量化失败: %s", doc.id, e)

    # 3. 存储
    if embeddings is not None and len(embeddings) == len(chunks):
        chunk_dicts = [
            {
                "index": c.index, "content": c.content, "embedding": emb,
                "token_count": c.token_count, "section_title": c.section_title,
            }
            for c, emb in zip(chunks, embeddings)
        ]
    else:
        # 向量化不可用，存储无向量的分块（后续可手动向量化）
        if embeddings is None:
            logger.info("向量化服务不可用，存储无向量分块: doc_id=%d", doc.id)
        chunk_dicts = [
            {
                "index": c.index, "content": c.content, "embedding": None,
                "token_count": c.token_count, "section_title": c.section_title,
            }
            for c in chunks
        ]
    vs = VectorStore(db)
    count = vs.store_chunks(doc.id, doc.kb_id, chunk_dicts)

    logger.info("文档 %d 自动索引完成: %d 个分块", doc.id, count)
    return count


def _refresh_kb_doc_count(db: Session, kb_id: int):
    cnt = db.execute(
        select(func.count(Document.id)).where(Document.kb_id == kb_id)
    ).scalar() or 0
    kb = db.get(KnowledgeBase, kb_id)
    if kb is not None:
        kb.doc_count = int(cnt)
        db.commit()
