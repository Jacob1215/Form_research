"""文档解析服务：优先 MinerU，失败降级到本地解析。"""
import os
import logging
from typing import Optional

import httpx

from .config import settings

logger = logging.getLogger("app.parsing")


def _read_text(path: str) -> str:
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return ""


def _parse_pdf_local(path: str) -> str:
    """使用 pdfplumber 解析 PDF（中文支持远优于 pypdf）。"""
    import pdfplumber
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                text = page.extract_text()
                if text:
                    parts.append(text)
            except Exception as e:  # noqa: BLE001
                logger.warning("PDF 第 %d 页解析失败：%s", i + 1, e)

    # 如果 pdfplumber 提取内容过少，尝试 PyMuPDF
    full_text = "\n".join(parts)
    if len(full_text.strip()) < 200:
        logger.info("PDF 可能为扫描件，尝试 PyMuPDF: %s", path)
        try:
            import fitz
            fitz_doc = fitz.open(path)
            ocr_parts = []
            for page in fitz_doc:
                text = page.get_text()
                if text and len(text.strip()) > 50:
                    ocr_parts.append(text)
            fitz_doc.close()
            if ocr_parts:
                return "\n".join(ocr_parts)
        except Exception as e:
            logger.warning("PyMuPDF 提取失败: %s", e)

    return full_text


def _parse_docx_local(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _parse_with_mineru(path: str) -> Optional[str]:
    """调用 MinerU POST /file_parse，返回 Markdown 文本；失败返回 None。"""
    if not settings.MINERU_API_URL:
        return None
    url = settings.MINERU_API_URL.rstrip("/") + "/file_parse"
    try:
        with httpx.Client(timeout=httpx.Timeout(180.0)) as client:
            with open(path, "rb") as f:
                files = {"file": (os.path.basename(path), f)}
                resp = client.post(url, files=files)
        if resp.status_code >= 400:
            logger.warning("MinerU 返回 %s: %s", resp.status_code, resp.text[:300])
            return None
        # 兼容多种返回结构
        ctype = resp.headers.get("content-type", "")
        if "application/json" in ctype:
            data = resp.json()
            # 常见字段：markdown / text / content / result
            for key in ("markdown", "text", "content", "result"):
                val = data.get(key)
                if isinstance(val, str) and val:
                    return val
                if isinstance(val, dict):
                    for sub in ("markdown", "text", "content"):
                        sub_val = val.get(sub)
                        if isinstance(sub_val, str) and sub_val:
                            return sub_val
            logger.warning("MinerU 返回 JSON 但无法识别内容字段：%s", str(data)[:300])
            return None
        # 纯文本/markdown 直接返回
        return resp.text
    except Exception as e:  # noqa: BLE001
        logger.warning("MinerU 调用异常：%s，降级本地解析。", e)
        return None


def parse_file(file_path: str, file_type: str) -> str:
    """解析文档，返回纯文本。file_type 不含点号。"""
    ext = (file_type or "").lower().lstrip(".")
    # 先尝试 MinerU（仅对 pdf/doc/docx 有意义）
    if ext in ("pdf", "doc", "docx"):
        md = _parse_with_mineru(file_path)
        if md:
            return md
    # 本地降级解析
    if ext in ("txt", "md"):
        return _read_text(file_path)
    if ext == "pdf":
        return _parse_pdf_local(file_path)
    if ext == "docx":
        return _parse_docx_local(file_path)
    if ext == "doc":
        # 老 .doc 二进制格式本地不解析，避免崩溃
        logger.warning("暂不支持本地解析 .doc 文件：%s", file_path)
        return ""
    # 兜底：按文本读取
    try:
        return _read_text(file_path)
    except Exception as e:  # noqa: BLE001
        logger.error("无法解析文件 %s：%s", file_path, e)
        return ""
